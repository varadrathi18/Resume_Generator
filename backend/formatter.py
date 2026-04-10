"""
formatter.py — Convert Gemini resume text to HTML, PDF, and DOCX.

Features:
  • Professional PDF with header bar, proper typography, and color scheme
  • Styled DOCX with proper heading hierarchy
  • Clean HTML with print-optimized CSS
  • Name-based file naming for easy classification
"""

import os
import re
import unicodedata
from datetime import datetime

import markdown
from docx import Document
from docx.shared import Pt, Inches, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from docx.oxml.ns import qn
from reportlab.lib.pagesizes import A4
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.lib.units import inch
from reportlab.lib.colors import HexColor
from reportlab.platypus import (
    SimpleDocTemplate, Paragraph, Spacer, HRFlowable,
)

from config import Config

GENERATED_DIR = Config.GENERATED_DIR
os.makedirs(GENERATED_DIR, exist_ok=True)


# ── Utility ─────────────────────────────────────────────────────────────

def _strip_md_preamble(text: str) -> str:
    """Remove any LLM preamble lines before the first '# ' heading."""
    lines = text.split("\n")
    for i, line in enumerate(lines):
        if line.strip().startswith("# "):
            return "\n".join(lines[i:])
    return text


def _sanitize_name(name: str) -> str:
    """
    Sanitize a person's name for use in filenames.
    Removes special characters, replaces spaces with underscores.
    Example: 'Varad Sharma' -> 'Varad_Sharma'
    """
    if not name or not name.strip():
        return "unnamed"

    # Normalize unicode characters
    name = unicodedata.normalize("NFKD", name)
    # Keep only alphanumeric and spaces
    name = re.sub(r"[^a-zA-Z0-9\s]", "", name)
    # Replace spaces with underscores
    name = name.strip().replace(" ", "_")
    # Remove consecutive underscores
    name = re.sub(r"_+", "_", name)
    # Limit length
    name = name[:50]

    return name or "unnamed"


def _generate_filename(name: str, extension: str) -> str:
    """
    Generate a descriptive filename using the person's name and timestamp.
    Example: 'Varad_Sharma_resume_2026-04-09_153022.pdf'
    """
    sanitized = _sanitize_name(name)
    timestamp = datetime.now().strftime("%Y-%m-%d_%H%M%S")
    return f"{sanitized}_resume_{timestamp}.{extension}"


# ── HTML ────────────────────────────────────────────────────────────────

def to_html(resume_text: str) -> str:
    """Convert markdown-formatted resume text to styled HTML."""
    cleaned = _strip_md_preamble(resume_text)
    body = markdown.markdown(cleaned, extensions=["extra", "sane_lists"])
    return f"""<!DOCTYPE html>
<html lang="en">
<head>
<meta charset="UTF-8"/>
<style>
  body {{
    font-family: 'Segoe UI', 'Inter', Tahoma, Geneva, Verdana, sans-serif;
    max-width: 800px; margin: 40px auto; padding: 0 24px;
    color: #1a1a2e; line-height: 1.65; font-size: 14px;
  }}
  h1 {{
    color: #0f172a; font-size: 28px; font-weight: 800;
    border-bottom: 3px solid #6366f1; padding-bottom: 10px;
    margin-bottom: 4px; letter-spacing: -0.5px;
  }}
  h1 + p {{ color: #475569; font-size: 13px; margin-top: 2px; margin-bottom: 20px; }}
  h2 {{
    color: #1e293b; font-size: 16px; font-weight: 700;
    border-bottom: 2px solid #e2e8f0; padding-bottom: 4px;
    margin-top: 24px; margin-bottom: 10px;
    text-transform: uppercase; letter-spacing: 0.5px;
  }}
  h3 {{
    color: #334155; font-size: 14px; font-weight: 600;
    margin-top: 14px; margin-bottom: 4px;
  }}
  h3 + p {{ color: #64748b; font-size: 12px; font-style: italic; margin-top: 0; }}
  ul {{ padding-left: 20px; margin-bottom: 8px; }}
  li {{
    margin-bottom: 5px; color: #334155; line-height: 1.65;
    font-size: 13.5px;
  }}
  p {{ color: #334155; line-height: 1.7; margin-bottom: 6px; font-size: 13.5px; }}
  strong {{ color: #0f172a; font-weight: 600; }}
  em {{ color: #64748b; }}
  @media print {{
    body {{ margin: 0; padding: 12px; font-size: 11px; }}
    h1 {{ font-size: 22px; border-bottom-width: 2px; }}
    h2 {{ font-size: 13px; margin-top: 16px; }}
  }}
</style>
</head>
<body>
{body}
</body>
</html>"""


# ── PDF ─────────────────────────────────────────────────────────────────

def _format_for_reportlab(text: str) -> str:
    """Convert markdown bold/italic markers to ReportLab-safe HTML tags."""
    # Convert bold
    text = re.sub(r"\*\*(.+?)\*\*", r"<b>\1</b>", text)
    text = re.sub(r"__(.+?)__", r"<b>\1</b>", text)
    # Convert italic
    text = re.sub(r"\*(.+?)\*", r"<i>\1</i>", text)
    text = re.sub(r"_(.+?)_", r"<i>\1</i>", text)
    # Escape ampersands for XML
    text = text.replace("&", "&amp;")
    # Fix double-escaped
    text = text.replace("&amp;amp;", "&amp;")
    return text.strip()


def to_pdf(resume_text: str, name: str = "") -> str:
    """Generate a professionally styled PDF from resume markdown text."""
    filename = _generate_filename(name, "pdf")
    filepath = os.path.join(GENERATED_DIR, filename)

    cleaned = _strip_md_preamble(resume_text)

    doc = SimpleDocTemplate(
        filepath,
        pagesize=A4,
        topMargin=0.5 * inch,
        bottomMargin=0.5 * inch,
        leftMargin=0.7 * inch,
        rightMargin=0.7 * inch,
    )

    styles = getSampleStyleSheet()

    # Custom styles
    PRIMARY = HexColor("#1e293b")
    ACCENT = HexColor("#6366f1")
    BODY_COLOR = HexColor("#334155")
    MUTED = HexColor("#64748b")

    styles.add(ParagraphStyle(
        "ResumeName",
        parent=styles["Heading1"],
        fontSize=22, spaceAfter=2, spaceBefore=0,
        textColor=PRIMARY, fontName="Helvetica-Bold",
        leading=26,
    ))
    styles.add(ParagraphStyle(
        "ContactInfo",
        parent=styles["Normal"],
        fontSize=10, spaceAfter=8, spaceBefore=0,
        textColor=MUTED, fontName="Helvetica",
        leading=14, alignment=0,
    ))
    styles.add(ParagraphStyle(
        "SectionHeading",
        parent=styles["Heading2"],
        fontSize=12, spaceAfter=4, spaceBefore=16,
        textColor=PRIMARY, fontName="Helvetica-Bold",
        borderPadding=(0, 0, 2, 0), leading=16,
    ))
    styles.add(ParagraphStyle(
        "SubHeading",
        parent=styles["Normal"],
        fontSize=11, spaceAfter=2, spaceBefore=8,
        textColor=HexColor("#1e293b"), fontName="Helvetica-Bold",
        leading=14,
    ))
    styles.add(ParagraphStyle(
        "DateLine",
        parent=styles["Normal"],
        fontSize=9.5, spaceAfter=4, spaceBefore=0,
        textColor=MUTED, fontName="Helvetica-Oblique",
        leading=12,
    ))
    styles.add(ParagraphStyle(
        "ResumeBody",
        parent=styles["Normal"],
        fontSize=10, leading=14, spaceAfter=3,
        textColor=BODY_COLOR, fontName="Helvetica",
    ))
    styles.add(ParagraphStyle(
        "BulletItem",
        parent=styles["Normal"],
        fontSize=10, leading=14, leftIndent=16, spaceAfter=3,
        bulletIndent=4, textColor=BODY_COLOR, fontName="Helvetica",
    ))

    elements: list = []
    lines = cleaned.split("\n")
    is_first_heading = True

    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            elements.append(Spacer(1, 4))
            continue

        if stripped.startswith("# "):
            text = _format_for_reportlab(stripped[2:])
            elements.append(Paragraph(text, styles["ResumeName"]))
            # Add accent line under the name
            elements.append(HRFlowable(
                width="100%", thickness=2, color=ACCENT,
                spaceBefore=2, spaceAfter=4,
            ))
            is_first_heading = False

        elif stripped.startswith("## "):
            text = _format_for_reportlab(stripped[3:]).upper()
            if not is_first_heading:
                elements.append(Spacer(1, 6))
            elements.append(Paragraph(text, styles["SectionHeading"]))
            # Subtle section divider
            elements.append(HRFlowable(
                width="100%", thickness=0.5, color=HexColor("#e2e8f0"),
                spaceBefore=0, spaceAfter=4,
            ))

        elif stripped.startswith("### "):
            text = _format_for_reportlab(stripped[4:])
            elements.append(Paragraph(text, styles["SubHeading"]))

        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            # Italic date line like *Jan 2023 – Present*
            text = _format_for_reportlab(stripped)
            elements.append(Paragraph(text, styles["DateLine"]))

        elif stripped.startswith(("- ", "* ", "• ")):
            text = _format_for_reportlab(stripped[2:])
            elements.append(Paragraph(f"•  {text}", styles["BulletItem"]))

        else:
            text = _format_for_reportlab(stripped)
            # Check if this is the contact line (right after name)
            if i <= 3 and ("|" in stripped or "@" in stripped):
                elements.append(Paragraph(text, styles["ContactInfo"]))
            else:
                elements.append(Paragraph(text, styles["ResumeBody"]))

    doc.build(elements)
    return filepath


# ── DOCX ────────────────────────────────────────────────────────────────

def _add_docx_runs(paragraph, text: str):
    """Parse text with **bold** and *italic* markers and add as runs."""
    # First handle bold
    parts = re.split(r'(\*\*.+?\*\*|__.+?__)', text)
    for part in parts:
        if not part:
            continue
        if (part.startswith('**') and part.endswith('**')) or \
           (part.startswith('__') and part.endswith('__')):
            content = part[2:-2]
            content = re.sub(r'[\*_]{1,2}(.*?)[*_]{1,2}', r'\1', content)
            run = paragraph.add_run(content)
            run.bold = True
        else:
            # Handle italic within non-bold text
            italic_parts = re.split(r'(\*.+?\*|_.+?_)', part)
            for ip in italic_parts:
                if not ip:
                    continue
                if (ip.startswith('*') and ip.endswith('*') and len(ip) > 2) or \
                   (ip.startswith('_') and ip.endswith('_') and len(ip) > 2):
                    run = paragraph.add_run(ip[1:-1])
                    run.italic = True
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)
                else:
                    clean_part = re.sub(r'[\*_]{1,2}(.*?)[*_]{1,2}', r'\1', ip)
                    paragraph.add_run(clean_part)


def _add_bottom_border(paragraph):
    """Add a subtle bottom border to a paragraph (section heading)."""
    pPr = paragraph.paragraph_format.element.get_or_add_pPr()
    pBdr = pPr.makeelement(qn('w:pBdr'), {})
    bottom = pBdr.makeelement(qn('w:bottom'), {
        qn('w:val'): 'single',
        qn('w:sz'): '4',
        qn('w:space'): '1',
        qn('w:color'): 'CBD5E1',
    })
    pBdr.append(bottom)
    pPr.append(pBdr)


def to_docx(resume_text: str, name: str = "") -> str:
    """Generate a professionally styled DOCX from resume markdown text."""
    filename = _generate_filename(name, "docx")
    filepath = os.path.join(GENERATED_DIR, filename)

    cleaned = _strip_md_preamble(resume_text)

    doc = Document()

    # Page margins
    for section in doc.sections:
        section.top_margin = Inches(0.5)
        section.bottom_margin = Inches(0.5)
        section.left_margin = Inches(0.7)
        section.right_margin = Inches(0.7)

    # Default style
    style = doc.styles["Normal"]
    font = style.font
    font.name = "Calibri"
    font.size = Pt(10.5)
    font.color.rgb = RGBColor(0x33, 0x41, 0x55)

    lines = cleaned.split("\n")
    for i, line in enumerate(lines):
        stripped = line.strip()
        if not stripped:
            doc.add_paragraph("")
            continue

        if stripped.startswith("# "):
            p = doc.add_heading('', level=0)
            _add_docx_runs(p, stripped[2:])
            p.alignment = WD_ALIGN_PARAGRAPH.LEFT
            for run in p.runs:
                run.font.size = Pt(22)
                run.font.color.rgb = RGBColor(0x0F, 0x17, 0x2A)

        elif stripped.startswith("## "):
            p = doc.add_heading('', level=1)
            text = stripped[3:].upper()
            _add_docx_runs(p, text)
            for run in p.runs:
                run.font.size = Pt(12)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)
            _add_bottom_border(p)

        elif stripped.startswith("### "):
            p = doc.add_heading('', level=2)
            _add_docx_runs(p, stripped[4:])
            for run in p.runs:
                run.font.size = Pt(11)
                run.font.color.rgb = RGBColor(0x1E, 0x29, 0x3B)

        elif stripped.startswith(("- ", "* ", "• ")):
            p = doc.add_paragraph(style="List Bullet")
            _add_docx_runs(p, stripped[2:])
            for run in p.runs:
                run.font.size = Pt(10.5)

        elif stripped.startswith("*") and stripped.endswith("*") and not stripped.startswith("**"):
            # Italic date line
            p = doc.add_paragraph()
            run = p.add_run(stripped.strip("*"))
            run.italic = True
            run.font.size = Pt(9.5)
            run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

        else:
            p = doc.add_paragraph()
            _add_docx_runs(p, stripped)
            # Contact line styling
            if i <= 3 and ("|" in stripped or "@" in stripped):
                for run in p.runs:
                    run.font.size = Pt(10)
                    run.font.color.rgb = RGBColor(0x64, 0x74, 0x8B)

    doc.save(filepath)
    return filepath
